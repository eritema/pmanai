import tkinter as tk
from tkinter import filedialog, messagebox
from api_handler.openai_integration import (
    abstract_summary_extraction,
    key_points_extraction,
    risks_extraction,
    action_item_extraction,
    sentiment_analysis,
    transcribe_audio,
    diagram_extraction
)
from docx import Document

class MainWindow(tk.Frame):
    def __init__(self, master=None):
        super(MainWindow, self).__init__(master)
        self.master = master
        self.setup_layout()
        self.create_menu()
        self.pack()
        self.file_entry= ""
        self.output_entry = ""
        # Layout configuration
        master.title("PManAI - Project Management AI Assistant")


    def create_menu(self):
        self.menu_bar = tk.Menu(self.master)
        self.master.config(menu=self.menu_bar)

        # File menu
        file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Open Audio File", command=self.open_audio_file)
        file_menu.add_separator()
        file_menu.add_command(label="Save as DOCX", command=lambda: self.save_file('docx'))
        file_menu.add_command(label="Save as MD", command=lambda: self.save_file('md'))


    def open_audio_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.mp3 *.wav *.ogg *.mp4")])
        if file_path:
            self.file_entry = file_path
            # Update any widgets if necessary to reflect the loaded data

    def save_file(self, file_type):

        file_path = filedialog.asksaveasfilename(defaultextension=f".{file_type}",
                                                 filetypes=[(f"{file_type.upper()} files", f"*.{file_type}")])
        if file_path:
            self.output_entry = file_path
            if file_type == 'docx':
                self.type = file_type
            elif file_type == 'md':
                self.type = file_type


    def setup_layout(self):
        # Add a Text widget
        self.transcription_text = tk.Text(self.master, height=10, width=50)
        self.transcription_text.pack()
        # Checkboxes for meeting minute components
        self.checks = {}
        for comp in ["Abstract Summary", "Key Points", "Risks", "Action Items", "Sentiment", "Diagram"]:
            self.checks[comp] = tk.BooleanVar()
            tk.Checkbutton(self.master, text=comp, variable=self.checks[comp]).pack()
        # Process audio button
        self.process_audio_button = tk.Button(self.master, text="Process audio", command=self.process_audio)
        self.process_audio_button.pack()

        # Process minutes button
        self.process_minutes_button = tk.Button(self.master, text="Process minutes", command=self.process_minutes)
        self.process_minutes_button.pack()
        pass

    def browse_file(self):
        filename = filedialog.askopenfilename()
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, filename)

    def process_audio(self):
        if not self.file_entry:
            messagebox.showwarning("Warning", "Please select an audio file and specify an output file name.")
            return
        audio_file_path = self.file_entry

        try:
            print(audio_file_path)
            print("Generating Trancription")
            
            self.transcription = transcribe_audio(audio_file_path)
            self.transcription_text.delete(1.0, tk.END)
            self.transcription_text.insert(tk.END, self.transcription)
            
            self.edited_transcription = self.transcription_text.get(1.0, tk.END).strip()
            messagebox.showinfo("Success", f"Transcribed text displayed in text field")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")


    def process_minutes(self):
        if not self.output_entry:
            messagebox.showwarning("Warning", "Please select an audio file and specify an output file name.")
            return
        output_file = self.output_entry
        print("Generating minutes")
        try:
            print(output_file)
            minutes = self.generate_meeting_minutes(self.edited_transcription)
            if self.type == 'docx':
                print("Word", f"Meeting minutes will be saved to {output_file}")
                self.save_as_docx(minutes, output_file)
            elif self.type == 'md':
                self.save_as_md(minutes, output_file)
                print("MD", f"Meeting minutes will be saved to {output_file}")
            messagebox.showinfo("Success", f"Meeting minutes saved to {output_file}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    def generate_meeting_minutes(self, transcription):
        minutes = {}
        if self.checks["Abstract Summary"].get():
            print("summary")
            minutes['abstract_summary'] = abstract_summary_extraction(transcription)
        if self.checks["Key Points"].get():
            print("Key points")
            minutes['key_points'] = key_points_extraction(transcription)
        if self.checks["Risks"].get():
            print("Risks")
            minutes['risks'] = risks_extraction(transcription)
        if self.checks["Action Items"].get():
            print("Items")
            minutes['action_items'] = action_item_extraction(transcription)
        if self.checks["Sentiment"].get():
            print("Sentiment")
            minutes['sentiment'] = sentiment_analysis(transcription)
        if self.checks["Diagram"].get():
            print("diagram")
            minutes['diagram'] = diagram_extraction(transcription)
        return minutes

    def save_as_docx(self, minutes, filename):
        doc = Document()
        for key, value in minutes.items():
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1)
            doc.add_paragraph(value)
            doc.add_paragraph()
        doc.save(filename)


    def save_as_md(self, minutes, filename):
        with open(filename, 'w', encoding='utf-8') as md_file:
            for key, value in minutes.items():
                heading = ' '.join(word.capitalize() for word in key.split('_'))
                md_file.write(f"## {heading}\n\n")  # Markdown heading level 2
                md_file.write(f"{value}\n\n")       # Paragraph
